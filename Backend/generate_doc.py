from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ── Helper functions ──────────────────────────────────────────────────────────
def set_font(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=(31, 73, 125))
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=(0, 112, 192))
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=(68, 68, 68))
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = val
            for para in cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph()

def divider():
    doc.add_paragraph("─" * 80)

# ════════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(30)
r = title_p.add_run("TravelMate AI")
r.font.size  = Pt(28)
r.font.bold  = True
r.font.color.rgb = RGBColor(31, 73, 125)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run("Backend — Complete Technical Summary")
r2.font.size  = Pt(16)
r2.font.color.rgb = RGBColor(0, 112, 192)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = date_p.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}   |   Status: COMPLETED")
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PROJECT OVERVIEW
# ════════════════════════════════════════════════════════════════════════════════
heading1("1. Project Overview")
body("TravelMate AI is a production-ready AI-powered travel planning backend built with FastAPI. "
     "It provides a complete REST API for user authentication, AI-generated trip itineraries, "
     "smart budget analytics, packing checklist management, and an AI concierge chat assistant.")

heading2("Technology Stack")
add_table(
    ["Component", "Technology", "Version"],
    [
        ["Web Framework",       "FastAPI",             ">=0.115.0"],
        ["ASGI Server",         "Uvicorn",             ">=0.30.1"],
        ["ORM",                 "SQLAlchemy",          ">=2.0.31"],
        ["Database",            "SQLite (dev)",        "—"],
        ["Migrations",          "Alembic",             ">=1.13.1"],
        ["Data Validation",     "Pydantic v2",         ">=2.9.0"],
        ["Authentication",      "JWT (python-jose)",   ">=3.3.0"],
        ["Password Hashing",    "passlib + bcrypt",    "4.0.1 pinned"],
        ["AI Provider",         "OpenAI / Gemini / Mock", "gpt-4o-mini / gemini-1.5-flash"],
        ["HTTP Client",         "httpx",               ">=0.27.0"],
        ["Settings",            "pydantic-settings",   ">=2.5.0"],
    ]
)

heading2("Architecture Pattern")
body("The project follows a clean layered architecture with strict separation of concerns:")
bullet("api/         — Thin FastAPI routers (request/response only)")
bullet("services/    — Business logic layer")
bullet("models/      — SQLAlchemy ORM models")
bullet("schemas/     — Pydantic request/response schemas")
bullet("core/        — Config, security utilities")
bullet("dependencies/— Reusable FastAPI dependencies (DB session, auth)")
bullet("prompts/     — AI prompt templates")
bullet("alembic/     — Database migration scripts")

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 2 — DATABASE DESIGN
# ════════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("2. Database Design")
body("The database consists of 5 tables with proper foreign key relationships and CASCADE deletes "
     "so removing a user or trip automatically cleans up all child records.")

heading2("Tables")
add_table(
    ["Table", "Primary Key", "Foreign Keys", "Purpose"],
    [
        ["users",         "id", "—",                            "User accounts"],
        ["trips",         "id", "user_id → users (CASCADE)",    "Trip records + AI itinerary JSON"],
        ["expenses",      "id", "trip_id → trips (CASCADE)",    "Expense logs per trip"],
        ["packing_items", "id", "trip_id → trips (CASCADE)",    "Packing checklist items"],
        ["chat_messages", "id", "trip_id → trips (CASCADE)",    "AI concierge chat history"],
    ]
)

heading2("Entity Relationships")
bullet("User  →  has many  →  Trips")
bullet("Trip  →  has many  →  Expenses")
bullet("Trip  →  has many  →  PackingItems")
bullet("Trip  →  has many  →  ChatMessages")

heading2("Alembic Migrations")
add_table(
    ["Revision ID", "Description", "Status"],
    [
        ["54fce337d14f", "Create initial tables (users, trips, expenses, packing_items)", "Applied"],
        ["8232e6700b89", "Add chat_messages table",                                        "Applied"],
    ]
)

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 3 — ALL API ENDPOINTS
# ════════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("3. Complete API Endpoint Reference")
body("All endpoints (except /health) are prefixed with /api. All protected endpoints require "
     "a JWT Bearer token in the Authorization header. Users can only access their own data — "
     "ownership is enforced on every endpoint.")

heading2("Authentication  —  /api/auth")
add_table(
    ["Method", "Endpoint", "Auth", "Description"],
    [
        ["POST", "/api/auth/register", "Public", "Register new user account"],
        ["POST", "/api/auth/login",    "Public", "Login and receive JWT access token"],
        ["GET",  "/api/auth/me",       "JWT",    "Get current logged-in user profile"],
    ]
)

heading2("Trips  —  /api/trips")
add_table(
    ["Method", "Endpoint", "Auth", "Description"],
    [
        ["POST",   "/api/trips/generate",  "JWT", "Generate AI trip itinerary + auto-populate expenses and packing"],
        ["GET",    "/api/trips/",          "JWT", "List all trips for current user"],
        ["GET",    "/api/trips/{id}",      "JWT", "Get a single trip by ID"],
        ["PUT",    "/api/trips/{id}",      "JWT", "Update trip fields"],
        ["DELETE", "/api/trips/{id}",      "JWT", "Delete trip and all cascade data"],
    ]
)

heading2("Budget Analytics  —  Expenses")
add_table(
    ["Method", "Endpoint", "Auth", "Description"],
    [
        ["GET",    "/api/trips/{id}/expenses",           "JWT", "List all expenses for a trip (sorted by date)"],
        ["POST",   "/api/trips/{id}/expenses",           "JWT", "Log a new expense against a trip"],
        ["GET",    "/api/trips/{id}/expenses/analytics", "JWT", "Budget breakdown by category with estimated vs actual totals"],
        ["GET",    "/api/expenses/{id}",                 "JWT", "Get a single expense record"],
        ["PUT",    "/api/expenses/{id}",                 "JWT", "Update an expense record"],
        ["DELETE", "/api/expenses/{id}",                 "JWT", "Delete an expense record"],
    ]
)

heading2("Packing Checklist")
add_table(
    ["Method", "Endpoint", "Auth", "Description"],
    [
        ["GET",    "/api/trips/{id}/packing",          "JWT", "List all packing items for a trip"],
        ["POST",   "/api/trips/{id}/packing",          "JWT", "Add a new packing item"],
        ["GET",    "/api/trips/{id}/packing/progress", "JWT", "Get packing progress summary and per-category stats"],
        ["GET",    "/api/packing/{id}",                "JWT", "Get a single packing item"],
        ["PUT",    "/api/packing/{id}",                "JWT", "Update a packing item"],
        ["PATCH",  "/api/packing/{id}/toggle",         "JWT", "Toggle packed/unpacked status"],
        ["DELETE", "/api/packing/{id}",                "JWT", "Delete a packing item"],
    ]
)

heading2("AI Concierge Chat")
add_table(
    ["Method", "Endpoint", "Auth", "Description"],
    [
        ["POST",   "/api/trips/{id}/chat",         "JWT", "Send message to AI concierge (context-aware, history-persistent)"],
        ["GET",    "/api/trips/{id}/chat/history", "JWT", "Retrieve full conversation history for a trip"],
        ["DELETE", "/api/trips/{id}/chat/history", "JWT", "Clear all chat messages for a trip"],
    ]
)

heading2("System")
add_table(
    ["Method", "Endpoint", "Auth", "Description"],
    [
        ["GET", "/health", "Public", "Server health check"],
        ["GET", "/docs",   "Public", "Swagger UI interactive documentation"],
        ["GET", "/redoc",  "Public", "ReDoc API documentation"],
    ]
)

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 4 — FILES CREATED AND MODIFIED
# ════════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("4. Files Created and Modified")

heading2("New Files Created")
add_table(
    ["File Path", "Purpose"],
    [
        ["app/api/chat.py",                                     "AI Concierge Chat API router (3 endpoints)"],
        ["app/models/chat.py",                                  "ChatMessage SQLAlchemy model"],
        ["app/schemas/chat.py",                                 "Pydantic schemas for chat request/response"],
        ["app/services/chat_service.py",                        "Chat business logic + AI provider routing"],
        ["app/services/expense_service.py",                     "Expense CRUD + analytics business logic"],
        ["app/services/packing_service.py",                     "Packing CRUD + progress calculation logic"],
        ["app/services/__init__.py",                            "Services package initializer"],
        ["alembic/versions/8232e6700b89_add_chat_messages.py",  "Alembic migration for chat_messages table"],
        ["tests/test_apis.py",                                  "Full 28-test API test suite (Python)"],
    ]
)

heading2("Modified Files")
add_table(
    ["File Path", "What Changed"],
    [
        ["app/api/expenses.py",     "Refactored to thin router; added GET /{id} and PUT /{id} endpoints; logic moved to service"],
        ["app/api/packing.py",      "Refactored to thin router; added GET /{id}, PUT /{id}, progress endpoints; logic moved to service"],
        ["app/main.py",             "Registered chat router"],
        ["app/models/trip.py",      "Added chat_messages relationship"],
        ["app/models/__init__.py",  "Added ChatMessage to imports"],
        ["app/prompts/templates.py","Added CHAT_SYSTEM_PROMPT_TEMPLATE and chat prompt constants"],
        ["requirements.txt",        "Pinned bcrypt==4.0.1; added requests>=2.32.0"],
        [".env.example",            "Added GEMINI_API_KEY example and AI_PROVIDER documentation"],
    ]
)

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 5 — STEP-BY-STEP IMPLEMENTATION SUMMARY
# ════════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("5. Implementation Summary")

heading2("Step 8 — Budget Analytics & Packing Checklist")
body("Objective: Complete expense and packing management with analytics and progress tracking.")

heading3("What Was Built:")
bullet("Expense Management")
bullet("Refactored expenses.py API router into a thin layer", level=1)
bullet("Created expense_service.py with all CRUD business logic", level=1)
bullet("Added GET /api/expenses/{id} — fetch single expense record", level=1)
bullet("Added PUT /api/expenses/{id} — update expense fields", level=1)
bullet("Enhanced analytics endpoint to include remaining_budget calculation", level=1)
bullet("")
bullet("Packing Checklist")
bullet("Refactored packing.py API router into a thin layer", level=1)
bullet("Created packing_service.py with all CRUD business logic", level=1)
bullet("Added GET /api/packing/{id} — fetch single packing item", level=1)
bullet("Added PUT /api/packing/{id} — full update of packing item fields", level=1)
bullet("Added GET /api/trips/{id}/packing/progress — progress summary", level=1)
bullet("Progress endpoint returns: total items, packed count, unpacked count, percentage, per-category breakdown", level=2)

heading3("Design Patterns Used:")
bullet("Service Layer Pattern — all business logic separated from routers")
bullet("DRY (Don't Repeat Yourself) — reusable helper functions for trip ownership checks")
bullet("Single Responsibility — routers only handle HTTP, services handle domain logic")

heading2("Step 9 — AI Concierge Chat")
body("Objective: Build a conversational AI assistant that knows about the user's trip context.")

heading3("What Was Built:")
bullet("Database & Models")
bullet("Created ChatMessage model with trip_id FK (CASCADE delete)", level=1)
bullet("Alembic migration 8232e6700b89 to create chat_messages table", level=1)
bullet("Updated Trip model with chat_messages relationship", level=1)
bullet("")
bullet("Schemas (Pydantic)")
bullet("ChatMessageRequest — validates user input (non-empty message)", level=1)
bullet("ChatMessageResponse — single message record", level=1)
bullet("ChatResponse — envelope containing both user + assistant messages", level=1)
bullet("ChatHistoryResponse — full conversation history array", level=1)
bullet("")
bullet("AI Service Layer")
bullet("chat_service.py — routes messages to OpenAI → Gemini → Mock fallback", level=1)
bullet("get_chat_history() — retrieves past messages ordered by time", level=1)
bullet("save_message() — persists user/assistant messages to DB", level=1)
bullet("clear_chat_history() — deletes all messages for a trip", level=1)
bullet("get_ai_response() — orchestrates history + system prompt + AI call", level=1)
bullet("_build_system_prompt() — enriches AI with trip context (destination, dates, budget, itinerary)", level=1)
bullet("_mock_response() — deterministic fallback when no API key configured", level=1)
bullet("")
bullet("AI Chat Router")
bullet("POST /api/trips/{id}/chat — send message, get AI reply", level=1)
bullet("GET /api/trips/{id}/chat/history — retrieve conversation", level=1)
bullet("DELETE /api/trips/{id}/chat/history — clear messages", level=1)
bullet("")
bullet("Prompt Engineering")
bullet("CHAT_SYSTEM_PROMPT_TEMPLATE — injects trip context (destination, dates, budget, itinerary)", level=1)
bullet("Includes last 20 messages of history to maintain conversation context", level=1)
bullet("AI is instructed to provide concise travel advice and stay in-domain", level=1)

heading3("AI Provider Support:")
bullet("OpenAI — gpt-4o-mini via API")
bullet("Gemini — gemini-1.5-flash via Google Generative AI API")
bullet("Mock — built-in deterministic responses (no API key required)")
bullet("Fallback chain: if primary provider fails → falls back to mock")

heading3("Security & Data Ownership:")
bullet("All chat endpoints are JWT-protected")
bullet("Ownership check: user can only access chat for trips they own")
bullet("Chat history isolated per trip")

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 6 — STEPS 1-7 (PRE-EXISTING)
# ════════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("6. Previously Completed (Steps 1-7)")
body("The following foundation was already in place before Steps 8 and 9 were implemented.")

heading2("Step 1-2 — Project Setup & Database Foundation")
bullet("FastAPI project structure with virtual environment")
bullet("SQLAlchemy engine, SessionLocal factory, declarative Base")
bullet("pydantic-settings for environment-based configuration")
bullet("SQLite database for development")
bullet("CORS middleware configured for React frontend (ports 3000, 5173)")

heading2("Step 3 — User Authentication (JWT)")
bullet("User model: id, email, hashed_password, full_name, is_active, is_admin, timestamps")
bullet("bcrypt password hashing via passlib")
bullet("JWT access token generation and validation via python-jose")
bullet("POST /api/auth/register — register with email + password")
bullet("POST /api/auth/login — login and receive Bearer token")
bullet("GET /api/auth/me — get current user profile (protected)")
bullet("get_current_user() dependency — reusable JWT guard for all protected routes")

heading2("Step 4 — Trip Management")
bullet("Trip model: title, destination, start/end dates, budget_level, traveler_type, interests (JSON), itinerary (JSON)")
bullet("Full CRUD: generate, list, get, update, delete")
bullet("Ownership enforcement on every route")
bullet("CASCADE delete removes expenses + packing + chat on trip deletion")

heading2("Step 5-6 — AI Integration (Itinerary Generation)")
bullet("AIProviderService — routes to OpenAI, Gemini, or Mock")
bullet("MockTravelPlannerService — full realistic itinerary without API keys")
bullet("ITINERARY_SYSTEM_PROMPT — strict JSON schema prompt for AI")
bullet("Auto-populates packing items and budget expenses from AI output on trip generation")
bullet("Fallback chain: API failure → Mock generator")

heading2("Step 7 — Alembic Migrations")
bullet("Migration 54fce337d14f creates all initial tables")
bullet("Full upgrade/downgrade support")
